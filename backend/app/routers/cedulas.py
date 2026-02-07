from fastapi import APIRouter, HTTPException
from supabase_client import supabase
from datetime import datetime
from fastapi import UploadFile, File
from fastapi.responses import JSONResponse
import uuid
from typing import List
import re
import unicodedata
import tempfile
import os
from fastapi.responses import FileResponse
from docx import Document
from tempfile import NamedTemporaryFile

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.shared import RGBColor


def limpiar_nombre_archivo(nombre: str) -> str:
    # Quitar acentos
    nombre = unicodedata.normalize("NFKD", nombre)
    nombre = nombre.encode("ascii", "ignore").decode("ascii")

    # Reemplazar espacios por _
    nombre = nombre.replace(" ", "_")

    # Eliminar caracteres raros
    nombre = re.sub(r"[^a-zA-Z0-9._-]", "", nombre)

    return nombre.lower()


router = APIRouter(prefix="/cedulas")

@router.get("/area/{area_id}")
def get_cedulas_by_area(area_id: int):
    return supabase.table("cedulas") \
        .select("id, criterio_nombre, estado") \
        .eq("area_id", area_id) \
        .order("id") \
        .execute().data


@router.get("/{id}")
def get_cedula(id: int):
    res = supabase.table("cedulas") \
        .select("*") \
        .eq("id", id) \
        .single() \
        .execute()

    if not res.data:
        raise HTTPException(status_code=404, detail="Cédula no encontrada")

    return res.data


@router.put("/{id}/estado")
def actualizar_estado(id: int, estado: int):
    return supabase.table("cedulas") \
        .update({"estado": estado}) \
        .eq("id", id) \
        .execute().data


@router.put("/{id}/valoracion")
def guardar_valoracion(id: int, data: dict):
    # 1️⃣ Obtener estado actual
    actual = supabase.table("cedulas") \
        .select("estado") \
        .eq("id", id) \
        .single() \
        .execute()

    if not actual.data:
        raise HTTPException(status_code=404, detail="Cédula no encontrada")

    estado_actual = actual.data["estado"]
    step = data.get("step")
    valoracion = data.get("valoracion")

    if step is None:
        raise HTTPException(status_code=400, detail="Step requerido")

    if not valoracion:
        raise HTTPException(status_code=400, detail="Valoración requerida")

    update_data = {
        "valoracion": valoracion,
        "fecha_envio": datetime.now().isoformat()
    }

    # 2️⃣ Avanzar estado SOLO si el step coincide
    if estado_actual == step:
        nuevo_estado = step + 1

        if 1 <= nuevo_estado <= 4:
            update_data["estado"] = nuevo_estado

    # 3️⃣ Ejecutar update
    res = supabase.table("cedulas") \
        .update(update_data) \
        .eq("id", id) \
        .execute()

    return {
        "ok": True,
        "actualizado": update_data
    }


@router.get("/codigo/{codigo}")
def get_cedulas_by_codigo(codigo: str):
    area_res = supabase.table("areas") \
        .select("id, nombre") \
        .eq("codigo", codigo) \
        .single() \
        .execute()

    if not area_res.data:
        raise HTTPException(status_code=404, detail="Área no encontrada")

    area = area_res.data

    cedulas = supabase.table("cedulas") \
        .select("id, estandar_nombre, estado, criterio_nombre") \
        .eq("area_id", area["id"]) \
        .order("id") \
        .execute()

    return {
        "area": area["nombre"],
        "cedulas": cedulas.data
    }

@router.put("/{id}/responsable")
def actualizar_responsable(id: int, data: dict):

    # 1️⃣ Obtener estado actual
    actual = supabase.table("cedulas") \
        .select("estado") \
        .eq("id", id) \
        .single() \
        .execute()

    if not actual.data:
        raise HTTPException(status_code=404, detail="Cédula no encontrada")

    estado_actual = actual.data["estado"]
    step = data.get("step")

    if step is None:
        raise HTTPException(status_code=400, detail="Step requerido")

    update_data = {}

    # 2️⃣ Actualizar responsable SOLO si vienen ambos campos
    nombre = data.get("nombre_responsable")
    cargo = data.get("cargo_responsable")

    if nombre is not None and cargo is not None:
        update_data.update({
            "nombre_responsable": nombre,
            "cargo_responsable": cargo
        })

    # 3️⃣ Avanzar estado SOLO si coincide el step
    if estado_actual == step:
        nuevo_estado = step + 1

        if 1 <= nuevo_estado <= 4:
            update_data["estado"] = nuevo_estado

    # 4️⃣ Evitar update vacío
    if not update_data:
        return {
            "ok": False,
            "mensaje": "No hubo cambios para actualizar"
        }

    # 5️⃣ Ejecutar update
    res = supabase.table("cedulas") \
        .update(update_data) \
        .eq("id", id) \
        .execute()

    return {
        "ok": True,
        "actualizado": update_data
    }

@router.post("/{cedula_id}/evidencias")
async def subir_evidencia(cedula_id: int, file: UploadFile = File(...)):
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="Solo PDFs")

    try:
        contenido = await file.read()
        nombre_limpio = limpiar_nombre_archivo(file.filename)
        filename = f"{uuid.uuid4()}_{nombre_limpio}"

        # 🔹 Subir al bucket público
        upload_res = supabase.storage.from_("evidencias").upload(
            filename,
            contenido,
            file_options={"content-type": "application/pdf"}
        )
        print("Upload response:", upload_res)
        if upload_res is None:
            raise HTTPException(status_code=500, detail="El archivo NO se subió al bucket. Revisa la API key y permisos.")

        # 🔹 Obtener URL pública permanente
        public_url = supabase.storage.from_("evidencias").get_public_url(filename)
        if not public_url:
            raise HTTPException(status_code=500, detail="No se pudo obtener la URL pública")

        # 🔹 Guardar en la DB
        res = supabase.table("evidencias").insert({
            "cedula_id": cedula_id,
            "nombre_archivo": file.filename,
            "archivo_url": public_url,
            "tipo": "PDF"
        }).execute()

        return {"ok": True, "archivo": res.data[0]}

    except Exception as e:
        print("❌ ERROR SUBIENDO PDF:", e)
        raise HTTPException(status_code=500, detail=str(e))

    
@router.get("/{cedula_id}/evidencias")
def listar_evidencias(cedula_id: int):
    res = supabase.table("evidencias") \
        .select("*") \
        .eq("cedula_id", cedula_id) \
        .order("id", desc=True) \
        .execute()

    return res.data

@router.delete("/evidencias/{evidencia_id}")
def eliminar_evidencia(evidencia_id: int):
    evidencia = supabase.table("evidencias") \
        .select("*") \
        .eq("id", evidencia_id) \
        .single() \
        .execute()

    if not evidencia.data:
        raise HTTPException(status_code=404, detail="Evidencia no encontrada")

    filename = evidencia.data["archivo_url"].split("/")[-1]

    supabase.storage.from_("evidencias").remove([filename])

    supabase.table("evidencias") \
        .delete() \
        .eq("id", evidencia_id) \
        .execute()

    return {"ok": True}

@router.get("/cedulasArea/{area_id}")
def get_cedulas_por_area_id(area_id: int):
    cedulas = supabase.table("cedulas") \
        .select("*") \
        .eq("area_id", area_id) \
        .order("id") \
        .execute()

    return {
        "cedulas": cedulas.data
    }


@router.get("/{cedula_id}/word")
def descargar_cedula_word(cedula_id: int):
    cedula = (
        supabase.table("cedulas")
        .select("*")
        .eq("id", cedula_id)
        .single()
        .execute()
        .data
    )

    if not cedula:
        raise HTTPException(status_code=404, detail="Cédula no encontrada")

    estandar = cedula["estandar_nombre"]
    criterio = cedula["criterio_nombre"]

    nombre_archivo = limpiar_filename(f"{estandar} {criterio}") + ".docx"

    doc = Document()

    # 🔹 TÍTULO DEL DOCUMENTO
    titulo = doc.add_heading(
        f"Cédula del {criterio} perteneciente al {estandar}",
        level=1,
    )
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    titulo.runs[0].font.name = "Montserrat"

    # 🔹 TABLA (1 columna)
    table = doc.add_table(rows=0, cols=1)
    table.style = "Table Grid"

    agregar_bloque(table, "Eje", cedula["eje_descripcion"])
    agregar_bloque(table, "Categoría", cedula["categoria_descripcion"])
    agregar_bloque(table, "Indicador", cedula["indicador_descripcion"])
    agregar_bloque(table, "Criterio", cedula["criterio_descripcion"])
    agregar_bloque(table, "Estándar", cedula["estandar_descripcion"])
    agregar_bloque(table, "Valoración argumentada", cedula["valoracion"])

    temp = NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp.name)
    temp.close()

    return FileResponse(
        temp.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=nombre_archivo,
    )

def set_cell_text(cell, text, *, bold=False, size=11, color=None, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text or "")
    run.bold = bold
    run.font.name = "Montserrat"
    run.font.size = Pt(size)

    if color:
        run.font.color.rgb = color

    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def titulo_fila(table, texto):
    row = table.add_row().cells
    row[0].merge(row[1])
    set_cell_text(
        row[0],
        texto,
        bold=True,
        size=13,
        color=RGBColor(0, 102, 204),
        align="center",
    )


def descripcion_fila(table, texto):
    row = table.add_row().cells
    row[0].merge(row[1])
    set_cell_text(
        row[0],
        texto,
        size=11,
        align="justify",
    )

def limpiar_filename(texto: str) -> str:
    texto = unicodedata.normalize("NFKD", texto)
    texto = texto.encode("ascii", "ignore").decode("ascii")
    texto = re.sub(r"[^\w\s-]", "", texto)
    texto = texto.strip().replace(" ", "_")
    return texto


def agregar_bloque(table, titulo, descripcion):
    row = table.add_row().cells
    cell = row[0]

    cell.text = ""
    p1 = cell.add_paragraph()
    run1 = p1.add_run(titulo)
    run1.bold = True
    run1.font.name = "Montserrat"
    run1.font.color.rgb = RGBColor(0, 102, 204)
    run1.font.size = Pt(12)
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p2 = cell.add_paragraph()
    run2 = p2.add_run(descripcion or "")
    run2.font.name = "Montserrat"
    run2.font.size = Pt(11)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
